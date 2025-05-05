# -*- coding: utf-8 -*-
"""
MedScribe AI Streamlit Application
Developed by Saqib Sherwani for Panacea Smart Solutions
"""


import streamlit as st
import google.generativeai as genai
import os
import re
import time
import io
import tempfile
from docx import Document
from PIL import Image
from fpdf import FPDF
from fpdf.enums import XPos, YPos
import html
import streamlit.components.v1 as components
import json

# --- Configuration ---
SI_FILENAME = "SI.txt"
APP_DIR = os.path.dirname(__file__) if "__file__" in locals() else os.getcwd()
SI_FILE_PATH = os.path.join(APP_DIR, SI_FILENAME)
ALLOWED_EXTENSIONS = [
    'mp3', 'wav', 'aac', 'ogg', 'flac',  # Audio
    'mp4', 'mpeg', 'mpg', 'mov', 'avi', 'flv', 'webm' # Video
]

# --- Helper Functions ---

@st.cache_data
def load_system_instruction(file_path):
    """Loads system instruction text from a file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            si_text = f.read()
        if not si_text:
            st.warning(f"System Instruction file ('{SI_FILENAME}') appears to be empty.", icon="⚠️")
            return None
        return si_text
    except FileNotFoundError:
        st.error(f"ERROR: System Instruction file ('{SI_FILENAME}') not found in the app directory ({APP_DIR}). Please ensure it exists.", icon="🚨")
        return None
    except Exception as e:
        st.error(f"ERROR: Failed to read System Instruction file: {e}", icon="🚨")
        return None

def configure_genai(api_key):
    """Configures the Google Generative AI client. Returns True on success, False otherwise."""
    if not api_key:
        return False
    try:
        genai.configure(api_key=api_key)
        return True
    except Exception as e:
        st.error(f"Error configuring Gemini API: {e}", icon="🚨")
        return False

def upload_to_gemini_api(tmp_file_path: str, mime_type: str):
    """Uploads a file (from a temporary local path) to the Gemini File API."""
    try:
        st.info(f"Uploading '{os.path.basename(tmp_file_path)}' to Gemini API...", icon="⏳")
        uploaded_file_resource = genai.upload_file(path=tmp_file_path, mime_type=mime_type)
        with st.spinner("Waiting for Gemini API file processing..."):
            while uploaded_file_resource.state.name == "PROCESSING":
                time.sleep(5)
                uploaded_file_resource = genai.get_file(uploaded_file_resource.name)
        if uploaded_file_resource.state.name == "FAILED":
            st.error(f"Gemini API File processing failed.", icon="🚨")
            try:
                if hasattr(uploaded_file_resource, 'error'): st.error(f"Failure details: {uploaded_file_resource.error}")
            except Exception: pass
            return None
        st.success(f"File '{uploaded_file_resource.display_name}' ready on Gemini API.", icon="✅")
        return uploaded_file_resource
    except Exception as e:
        st.error(f"Error uploading file to Gemini API: {e}", icon="🚨")
        if "API_KEY_INVALID" in str(e): st.error("The provided API Key seems invalid. Please check and re-enter.", icon="🔑")
        return None

def generate_note_from_api(audio_resource, si_text):
    """Generates the SOAP note using the Gemini API. Returns raw text or None."""
    if not audio_resource: st.error("Cannot generate note: Invalid audio resource provided.", icon="🚨"); return None

    full_response_text = ""
    # --- UPDATED Model Name ---
    model_name = "gemini-2.0-flash" # Using the specific model requested

    generation_config = genai.types.GenerationConfig(temperature=1, top_p=1, max_output_tokens=8192)
    file_display_name = getattr(audio_resource, 'display_name', audio_resource.name)
    user_prompt_text = f"Please generate a detailed SOAP note in the specified format for the following provided audio recording ({file_display_name}) of a provider-patient interaction by carefully following the SI.txt."
    contents = [audio_resource, user_prompt_text]
    try:
        model = genai.GenerativeModel(model_name=model_name, generation_config=generation_config, system_instruction=si_text)
        with st.spinner(f"Generating SOAP note via {model_name} API..."): # Show model name in spinner
            response = model.generate_content(contents, stream=True)
            for chunk in response:
                try:
                    if hasattr(chunk, 'text') and chunk.text: full_response_text += chunk.text
                except Exception as chunk_e: st.warning(f"Warning: Error processing individual response chunk: {chunk_e}", icon="⚠️"); pass
        # Post-stream checks
        if not full_response_text and hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
             st.error(f"Generation blocked by API due to: {response.prompt_feedback.block_reason}", icon="🚫")
             return None
        if not full_response_text:
            # Try to get finish reason if available
            finish_reason = "Unknown"
            try:
                if hasattr(response, 'candidates') and response.candidates:
                    finish_reason = response.candidates[0].finish_reason.name
            except Exception: pass
            st.warning(f"Generation finished, but no text content was received. Finish reason: {finish_reason}", icon="⚠️")
            return None

        st.success("SOAP note generated successfully.", icon="✍️")
        return full_response_text
    except Exception as e:
        st.error(f"An error occurred during API call with {model_name}: {e}", icon="🚨")
        if "API_KEY_INVALID" in str(e): st.error("The provided API Key is invalid or expired.", icon="🔑")
        elif "RESOURCE_EXHAUSTED" in str(e): st.error("Quota possibly exceeded. Please check your Google Cloud project quotas.", icon="📊")
        elif "not found" in str(e).lower() and model_name in str(e): st.error(f"Model '{model_name}' not found or not available. Check model name or your API access.", icon="❌")
        return None

# Regex Parser (remains the same)
def parse_soap_note_regex(note_text: str) -> dict:
    """Parses the generated SOAP note string using robust regex. Cleans markdown/numbering."""
    header_text_map = { "CC": ["chief complaint", "cc"], "HPI": ["history of present illness", "hpi"], "PE": ["physical exam", "physical examination", "pe"], "Assessment": ["assessment", "impression"],"Plan": ["plan"]}
    text_to_key_map = {}; all_header_texts_list = []
    for key, texts in header_text_map.items():
        for text in texts: text_to_key_map[text] = key; all_header_texts_list.append(re.escape(text))
    output_keys = list(header_text_map.keys()); sections = {key: "" for key in output_keys}; sections["parsing_notes"] = []

    def clean_content(text):
        if not isinstance(text, str): return ""
        cleaned = text.replace('**', '').replace('*', ''); cleaned = cleaned.strip(': \n\r\t')
        cleaned = re.sub(r"^\s*\d+\.\s*", "", cleaned, flags=re.MULTILINE); cleaned = re.sub(r"\s*\d+\.\s*$", "", cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r'\n\s*\n', '\n', cleaned).strip(); return cleaned

    all_header_texts_pattern = "|".join(all_header_texts_list)
    header_pattern = re.compile(r"(?:^\s*\d+\.\s+)?(?:[*]{2})?" rf"({all_header_texts_pattern})" r"(?:\s*\([^)]+\))?(?:[*]{2})?" r"\s*:", re.IGNORECASE | re.MULTILINE)
    found_headers = []
    for match in header_pattern.finditer(note_text):
        matched_text_lower = match.group(1).lower(); key = text_to_key_map.get(matched_text_lower)
        if key:
            is_duplicate = False
            if found_headers: last_header = found_headers[-1]; is_duplicate = last_header["key"] == key and (match.start() - last_header["start"]) < 20
            if not is_duplicate: found_headers.append({"key": key, "start": match.start(), "end": match.end()})
    found_headers.sort(key=lambda x: x["start"])

    if not found_headers:
        sections["parsing_notes"].append("No standard SOAP headers found using regex pattern.")
        sections["HPI"] = clean_content(note_text)
        for key in output_keys:
             if key != "HPI": sections[key] = "[SECTION NOT FOUND IN RESPONSE]"
        final_sections_no_hdr = {k: v for k, v in sections.items() if k in output_keys};
        if sections.get("parsing_notes"): final_sections_no_hdr["parsing_notes"] = sections["parsing_notes"]
        return final_sections_no_hdr

    for i, current_header in enumerate(found_headers):
        key = current_header["key"]; start_content_idx = current_header["end"]; end_content_idx = len(note_text)
        if i + 1 < len(found_headers): end_content_idx = found_headers[i+1]["start"]
        content_raw = note_text[start_content_idx:end_content_idx]
        if not sections[key]: sections[key] = clean_content(content_raw)

    found_keys = {h["key"] for h in found_headers}
    for key in output_keys:
        if key not in found_keys:
             sections[key] = "[SECTION NOT FOUND IN RESPONSE]"; primary_term = header_text_map[key][0]
             sections["parsing_notes"].append(f"Content for section '{primary_term}' not found.")

    final_sections = {k: v for k, v in sections.items() if k in output_keys}
    if sections.get("parsing_notes"): final_sections["parsing_notes"] = sections["parsing_notes"]
    return final_sections


# --- Export Functions --- (Remain the same)
def create_docx_report(sections_dict):
    """Creates a DOCX document in memory reading current values from text areas."""
    document = Document()
    document.add_heading('MedScribe AI - Generated SOAP Note', 0)
    section_order = ["CC", "HPI", "PE", "Assessment", "Plan"]
    for key in section_order:
        content = st.session_state.get(f"edit_{key}", sections_dict.get(key, "[Not Provided or Found]"))
        document.add_heading(key, level=1)
        for para in content.split('\n'):
             if para.strip(): document.add_paragraph(para.strip())
        document.add_paragraph() # Space
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Define paths to the Noto Sans font files (assuming they are in the same directory)
APP_DIR = os.path.dirname(__file__) if "__file__" in locals() else os.getcwd()
# --- Make sure these filenames match the font files you actually have ---
NOTO_FONT_REGULAR_FILENAME = "NotoSans-Regular.ttf"
NOTO_FONT_BOLD_FILENAME = "NotoSans-Bold.ttf" # Optional bold version
NOTO_FONT_REGULAR_PATH = os.path.join(APP_DIR, NOTO_FONT_REGULAR_FILENAME)
NOTO_FONT_BOLD_PATH = os.path.join(APP_DIR, NOTO_FONT_BOLD_FILENAME)

# --- UPDATED create_pdf_report (Removed Toasts) ---
def create_pdf_report(sections_dict):
    """Creates a PDF document using embedded Noto Sans Unicode font."""
    pdf = FPDF()
    pdf.add_page()

    # --- Add Noto Sans Unicode Font ---
    font_loaded_regular = False
    font_loaded_bold = False
    try:
        # Add the regular font. REMOVED uni=True
        pdf.add_font('NotoSans', '', NOTO_FONT_REGULAR_PATH) # Assumes NOTO_FONT_REGULAR_PATH is defined above
        font_loaded_regular = True
        # REMOVED st.toast(...)

        # Attempt to add the bold font if it exists. REMOVED uni=True
        if os.path.exists(NOTO_FONT_BOLD_PATH): # Assumes NOTO_FONT_BOLD_PATH is defined above
           pdf.add_font('NotoSans', 'B', NOTO_FONT_BOLD_PATH)
           font_loaded_bold = True
           # REMOVED st.toast(...)
        # else: # Optional info toast removed as well
           # REMOVED st.toast(...)
           # pass # No action needed if bold doesn't exist

    except FileNotFoundError as fnf_error:
        missing_font_path = str(fnf_error).split("'")[-2]
        st.error(f"Font file not found: {missing_font_path}. Please ensure Noto Sans font files are in the script directory ({APP_DIR}). PDF generation failed.", icon="📄")
        return None
    except RuntimeError as e:
         st.error(f"FPDF Error adding Noto Sans font: {e}. Ensure the font files are valid.", icon="📄")
         return None

    if not font_loaded_regular:
        st.error("Regular Noto Sans font could not be loaded. Cannot generate PDF.", icon="❌")
        return None

    # --- Use the Noto Sans font throughout the document ---
    # Title
    pdf.set_font("NotoSans", size=16)
    pdf.cell(0, 10, "MedScribe AI - Generated SOAP Note", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.ln(10)

    section_order = ["CC", "HPI", "PE", "Assessment", "Plan"]
    current_section_key_for_error = "Unknown"
    try:
        for key in section_order:
            current_section_key_for_error = key
            content = st.session_state.get(f"edit_{key}", sections_dict.get(key, "[Not Provided or Found]"))

            # Section Header Font
            pdf.set_font("NotoSans", 'B', size=12)
            pdf.multi_cell(0, 10, key, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Section Content Font
            pdf.set_font("NotoSans", size=11)
            pdf.multi_cell(0, 7, content, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(5)

        # --- Generate PDF Bytes ---
        pdf_output = pdf.output()
        pdf_output_bytes = bytes(pdf_output)
        return pdf_output_bytes
    except Exception as e:
        st.error(f"Failed to generate PDF bytes (Section: {current_section_key_for_error}): {e}", icon="🚨")
        return None

# --- REVISED Clipboard Function using data attributes ---
def copy_button_html(text_to_copy: str, button_key: str, button_text: str = "📋", help_text: str = "Copy to clipboard"):
    """Generates HTML/JS component for a copy button using data attributes."""

    # Escape text safely for insertion into an HTML data attribute
    escaped_data_text = html.escape(text_to_copy, quote=True)

    # Unique function name for each button instance to avoid JS conflicts
    js_function_name = f"copyText_{button_key.replace('-', '_').replace('.', '_')}" # Ensure valid JS function name chars

    component_html = f"""
    <script>
        // Define the function that will handle the copy action
        function {js_function_name}(buttonElement) {{
            // Retrieve the text from the button's own data-text attribute
            const textToCopy = buttonElement.getAttribute('data-text');

            if (!navigator.clipboard) {{
                console.error("Clipboard API not available");
                buttonElement.innerText = 'Error!'; // Indicate API not available
                setTimeout(() => {{ buttonElement.innerText = '{html.escape(button_text)}'; }}, 2000);
                return;
            }}

            navigator.clipboard.writeText(textToCopy)
                .then(() => {{
                    const originalText = '{html.escape(button_text)}'; // Get original button text
                    buttonElement.innerText = 'Copied!';
                    buttonElement.disabled = true;
                    setTimeout(() => {{
                        // Check if element still exists before resetting
                        const currentBtn = document.getElementById(buttonElement.id);
                        if (currentBtn) {{
                           currentBtn.innerText = originalText;
                           currentBtn.disabled = false;
                        }}
                    }}, 1500); // Revert after 1.5 seconds
                }})
                .catch(err => {{
                    console.error('Failed to copy text for key {button_key}: ', err);
                    const originalText = '{html.escape(button_text)}';
                    buttonElement.innerText = 'Error';
                    buttonElement.disabled = true;
                    setTimeout(() => {{
                        const currentBtn = document.getElementById(buttonElement.id);
                         if (currentBtn) {{
                            currentBtn.innerText = originalText;
                            currentBtn.disabled = false;
                         }}
                    }}, 1500);
                }});
        }}
    </script>

    <button
        id="copyBtn_{button_key}"
        title="{html.escape(help_text)}"
        onclick="{js_function_name}(this); event.stopPropagation();"
        data-text="{escaped_data_text}"
        style="
            background: none; border: none; padding: 0.1rem 0.3rem; margin: 0;
            font-size: 0.9em; cursor: pointer; line-height: 1; vertical-align: middle;
            float: right; position: relative; top: 5px; right: 5px;
        "
    >
        {html.escape(button_text)} </button>
    """
    # Set scrolling=False; sometimes helps with component rendering stability
    return components.html(component_html, height=40, scrolling=False)

# --- Streamlit App UI ---

st.set_page_config(page_title="MedScribe AI", layout="wide")

# --- Sidebar --- (Remains the same)
with st.sidebar:
    st.markdown("## 🔑 Configuration")
    st.markdown("Enter your Gemini API Key below. Your key is used solely for processing your request during this session.")
    if 'api_key' not in st.session_state: st.session_state.api_key = ''
    api_key_input = st.text_input(
        "Gemini API Key", type="password", key="api_key_input_widget",
        value=st.session_state.api_key, help="Obtain your key from Google AI Studio"
    )
    if api_key_input != st.session_state.api_key:
        st.session_state.api_key = api_key_input
        st.session_state.raw_text = None; st.session_state.parsed_sections = None
        st.session_state.error_message = None; st.session_state.gemini_resource_name = None
        st.session_state.uploaded_file_info = None
        for key in ["CC", "HPI", "PE", "Assessment", "Plan"]:
            if f"edit_{key}" in st.session_state: del st.session_state[f"edit_{key}"]
        st.rerun()
    st.markdown("💡 **Tip:** Obtain your key from [Google AI Studio](https://aistudio.google.com/app/apikey)")
    st.markdown("---")
    st.markdown("<small>MedScribe AI © 2025</small>", unsafe_allow_html=True)

# --- Main Page ---
st.title("🩺 MedScribe AI")
st.caption("AI-Assisted Medical Documentation Tool - By Saqib Sherwani")

# Initialize session state variables
if 'raw_text' not in st.session_state: st.session_state.raw_text = None
if 'parsed_sections' not in st.session_state: st.session_state.parsed_sections = None
if 'error_message' not in st.session_state: st.session_state.error_message = None
if 'processing' not in st.session_state: st.session_state.processing = False
if 'gemini_resource_name' not in st.session_state: st.session_state.gemini_resource_name = None
if 'uploaded_file_info' not in st.session_state: st.session_state.uploaded_file_info = None

# Load System Instruction
system_instruction = load_system_instruction(SI_FILE_PATH)
if not system_instruction and 'si_load_error_shown' not in st.session_state:
     st.error("Failed to load system instructions. Please ensure SI.txt exists in the app folder.", icon="📄")
     st.session_state.si_load_error_shown = True

# --- File Uploader --- (Remains the same)
uploaded_file = st.file_uploader(
    "Upload Audio or Video File", type=ALLOWED_EXTENSIONS, accept_multiple_files=False,
    key="file_uploader_widget", help="Upload patient encounter recording (MP3, WAV, MP4, MOV, etc.)"
)
if uploaded_file is not None:
    if st.session_state.uploaded_file_info is None or st.session_state.uploaded_file_info['name'] != uploaded_file.name or st.session_state.uploaded_file_info['size'] != uploaded_file.size:
           st.session_state.uploaded_file_info = { 'name': uploaded_file.name, 'type': uploaded_file.type, 'size': uploaded_file.size, 'data': uploaded_file.getvalue() }
           st.session_state.raw_text = None; st.session_state.parsed_sections = None; st.session_state.error_message = None
           for key in ["CC", "HPI", "PE", "Assessment", "Plan"]:
               if f"edit_{key}" in st.session_state: del st.session_state[f"edit_{key}"]

# --- Generate Button and Processing Logic ---
col1, col2 = st.columns([1, 5])
with col1:
    button_disabled = not st.session_state.uploaded_file_info or not st.session_state.api_key or st.session_state.processing or not system_instruction
    generate_button = st.button("Generate SOAP Note", type="primary", disabled=button_disabled, use_container_width=True)

if generate_button and st.session_state.uploaded_file_info:
    # Reset state
    st.session_state.raw_text = None; st.session_state.parsed_sections = None
    st.session_state.error_message = None; st.session_state.processing = True
    st.session_state.gemini_resource_name = None # Ensure reset before try block
    temp_file_path = None; gemini_resource = None
    current_file_info = st.session_state.uploaded_file_info

    # Configure API
    if not configure_genai(st.session_state.api_key):
        st.session_state.error_message = "Gemini API Key configuration failed. Please check the key."
        st.session_state.processing = False; st.rerun()

    # Main processing block
    try:
        with st.spinner(f"Processing '{current_file_info['name']}'..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(current_file_info['name'])[1]) as tmp_file:
                tmp_file.write(current_file_info['data']); temp_file_path = tmp_file.name
            gemini_resource = upload_to_gemini_api(temp_file_path, current_file_info['type'])

            if gemini_resource:
                st.session_state.gemini_resource_name = gemini_resource.name # Set resource name *only* on successful upload
                raw_text_result = generate_note_from_api(gemini_resource, system_instruction)
                st.session_state.raw_text = raw_text_result
                if raw_text_result:
                    with st.spinner("Parsing generated text..."):
                        st.session_state.parsed_sections = parse_soap_note_regex(raw_text_result)
                        if st.session_state.parsed_sections:
                            for key in ["CC", "HPI", "PE", "Assessment", "Plan"]:
                                st.session_state[f"edit_{key}"] = st.session_state.parsed_sections.get(key, "")
                        else:
                             if not st.session_state.error_message: st.session_state.error_message = "Generated text, but failed to parse into SOAP sections."
                else:
                    # Error message is set within generate_note_from_api if needed
                    if not st.session_state.error_message: st.session_state.error_message = "Failed to generate text content from the API for unknown reasons."
            else:
                # Error message is set within upload_to_gemini_api if needed
                if not st.session_state.error_message: st.session_state.error_message = "Failed to upload file to Gemini API for unknown reasons."
    except Exception as e:
        st.session_state.error_message = f"An unexpected error occurred during processing: {e}"
        st.error(st.session_state.error_message, icon="🚨") # Show error immediately
    finally:
        # --- Cleanup ---
        if temp_file_path and os.path.exists(temp_file_path):
            try: os.remove(temp_file_path)
            except Exception as e: st.warning(f"Could not delete temporary file {temp_file_path}: {e}", icon="⚠️")

        # --- FIXED Cleanup Logic ---
        # Remove the genai.is_configured() check
        if st.session_state.gemini_resource_name: # Check if resource name exists (implies successful config/upload)
             try:
                 st.info(f"Attempting to clean up Gemini file: {st.session_state.gemini_resource_name}", icon="🧹") # More visible cleanup attempt message
                 genai.delete_file(st.session_state.gemini_resource_name)
                 st.toast(f"Cleaned up file {st.session_state.gemini_resource_name} from Gemini API.", icon="✅")
                 st.session_state.gemini_resource_name = None # Clear on success
             except Exception as e:
                 st.warning(f"Could not delete file {st.session_state.gemini_resource_name} from Gemini API: {e}. Manual cleanup might be needed via Google AI Studio.", icon="⚠️")

        st.session_state.processing = False
        st.rerun() # Update UI

# --- Display Results or Errors --- (UI part remains mostly the same)
st.markdown("---")

if st.session_state.processing:
    st.info("Processing... Please wait.", icon="⏳")

if st.session_state.error_message and not st.session_state.processing:
    st.error(st.session_state.error_message, icon="🚨")

if st.session_state.parsed_sections and not st.session_state.processing:
    st.subheader("Generated SOAP Note Sections")
    sections_data = st.session_state.parsed_sections

    # --- Export Buttons ---
    st.markdown("**Export Options:**")
    export_cols = st.columns(3)
    with export_cols[0]:
        try:
            docx_buffer = create_docx_report(sections_data or {})
            st.download_button(label="Export as DOCX", data=docx_buffer, file_name="MedScribeAI_SOAP_Note.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="docx_download")
        except Exception as e: st.warning(f"Could not generate DOCX: {e}", icon="⚠️")
    with export_cols[1]:
         try:
             pdf_bytes_output = create_pdf_report(sections_data or {})
             if pdf_bytes_output:
                 st.download_button(label="Export as PDF", data=pdf_bytes_output, file_name="MedScribeAI_SOAP_Note.pdf", mime="application/pdf", use_container_width=True, key="pdf_download")
         except Exception as e: st.warning(f"PDF generation/download failed: {e}", icon="⚠️")

    st.markdown("---")
    st.markdown("**Edit Sections:** (Changes are included in export)")

    # --- Display EDITABLE Sections with COPY BUTTON ---
    section_order = ["CC", "HPI", "PE", "Assessment", "Plan"]
    for key in section_order:
        st.subheader(key)
        st.text_area(
            label=f"edit_{key}_label",
            height=200 if len(st.session_state.get(f"edit_{key}", "")) > 150 else 100,
            key=f"edit_{key}",
            label_visibility="collapsed",
            help=f"Edit the {key} section here."
        )
        copy_button_html(
            text_to_copy=st.session_state.get(f"edit_{key}", ""),
            button_key=f"copy_{key}",
            button_text="📋",
            help_text=f"Copy {key} to clipboard"
        )
        st.write("")


    # --- Display Raw Output with COPY BUTTON ---
    if st.session_state.raw_text:
        st.markdown("---")
        with st.expander("Raw Gemini Output (for comparison)"):
            copy_button_html(
                text_to_copy=st.session_state.raw_text,
                button_key="copy_raw",
                button_text="📋 Raw",
                help_text="Copy raw output to clipboard"
            )
            st.text_area(
                "Raw Output",
                value=st.session_state.raw_text,
                height=300,
                disabled=True,
                key="raw_output_display",
                label_visibility="collapsed",
                help="This is the raw, unparsed text from the AI."
            )

    # Display parsing notes
    if sections_data and sections_data.get("parsing_notes"):
        with st.expander("Parsing Notes"):
             for note in sections_data["parsing_notes"]: st.caption(f"- {note}")

elif not st.session_state.processing and not st.session_state.error_message:
     st.info("Enter your API key in the sidebar, upload an audio/video file, and click 'Generate SOAP Note'.", icon="☝️")

# --- Footer ---
st.markdown("---")
st.markdown(
    """<div style="text-align: center; font-size: small; color: grey;">
    MedScribe AI © 2025 | Panacea Smart Solutions | AI-Assisted Medical Documentation Tool - Developed by Saqib Sherwani
    </div>""",
    unsafe_allow_html=True
)